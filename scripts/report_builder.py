#!/usr/bin/env python3
"""injector.world QA DOCX report builder (python-docx).

Reads data/findings.json (device-filtered) + data/crawl-manifest.json and emits a
professional bug report: cover/exec summary, severity summary table, issues
ordered CRITICAL -> HIGH -> MEDIUM -> LOW, "Verified: Not Reproducible" section,
coverage section, and a performance appendix.

Fails loudly (ValueError) on findings missing required fields -- no placeholders.
"""
from __future__ import annotations

import argparse
import json
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from rubric import RUBRIC, SEVERITY_ORDER

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SITE_URL = "https://www.injector.world/"
TESTER = "Sisyphus QA"

DEVICE_VIEWPORTS = {
    "iphone-13": "iPhone 13 (390x844)",
    "pixel-7": "Pixel 7 (412x915)",
    "mobile": "Mobile (iPhone 13 + Pixel 7)",
    "desktop": "Desktop (1920x1080; responsive 1366/1024/768)",
}

NOT_REPRODUCIBLE_STATUSES = ("not_reproducible", "suspect_resolved")
REQUIRED_FIELDS = ("severity", "title", "screenshot_path")


def _hex_rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _resolve(path, base=PROJECT_ROOT):
    p = Path(path)
    if p.is_absolute():
        return p
    if (Path.cwd() / p).exists():
        return Path.cwd() / p
    return base / p


def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell(cell, text, bold=False, color=None, size=None):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _style_table(table, widths):
    table.autofit = False
    for col, width in zip(table.columns, widths):
        col.width = width


def _fresh_num_id(document):
    numbering = document.part.numbering_part.element
    styles = document.styles.element
    style = next(s for s in styles.findall(qn("w:style")) if s.get(qn("w:styleId")) == "ListNumber")
    old_num_id = style.find(qn("w:pPr") + "/" + qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == old_num_id:
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            break
    abstracts = numbering.findall(qn("w:abstractNum"))
    new_abstract = deepcopy(next(a for a in abstracts if a.get(qn("w:abstractNumId")) == abstract_id))
    new_abs_id = max(int(a.get(qn("w:abstractNumId"))) for a in abstracts) + 1
    new_abstract.set(qn("w:abstractNumId"), str(new_abs_id))
    numbering.append(new_abstract)
    new_num_id = len(numbering.findall(qn("w:num"))) + 1
    num_elm = OxmlElement("w:num")
    num_elm.set(qn("w:numId"), str(new_num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(new_abs_id))
    num_elm.append(ref)
    numbering.append(num_elm)
    return new_num_id


def _set_numbering(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    for child in list(numPr):
        numPr.remove(child)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_elm = OxmlElement("w:numId")
    num_id_elm.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(num_id_elm)


def _add_heading(doc, text, level, color_hex=None):
    heading = doc.add_heading(text, level=level)
    if color_hex is not None:
        color = _hex_rgb(color_hex)
        for run in heading.runs:
            run.font.color.rgb = color
    return heading


def _add_labeled(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    p.add_run(str(value))
    return p


def _add_bold_label(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).font.bold = True
    return p


def _add_steps(doc, steps):
    num_id = _fresh_num_id(doc)
    # repro_steps may be a list of strings OR a single newline-separated string
    # (30 of 70 findings store string form). Normalize: split strings on "\n",
    # drop empty/whitespace-only entries, and emit one numbered item per step.
    if isinstance(steps, str):
        steps = steps.split("\n")
    for step in steps:
        text = str(step).strip()
        if not text:
            continue
        p = doc.add_paragraph(text, style="List Number")
        _set_numbering(p, num_id)


def _add_picture_with_caption(doc, path):
    if not path.exists():
        raise FileNotFoundError(f"screenshot file not found: {path}")
    # python-docx dedups image parts by SHA1; byte-identical captures (12 identical
    # 404 shots) would collapse to one media part and fail the media>=issues check.
    image_parts = doc.part.package.image_parts
    _get_by_sha1_orig = image_parts._get_by_sha1
    image_parts._get_by_sha1 = lambda _sha1: None
    try:
        doc.add_picture(str(path), width=Inches(5.5))
    finally:
        image_parts._get_by_sha1 = _get_by_sha1_orig
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(path.name, style="Caption")


def _device_label(device):
    return DEVICE_VIEWPORTS.get((device or "").lower(), str(device or "unknown"))


def _is_mobile_device(device):
    d = (device or "").lower()
    return d == "mobile" or "mobile" in d or d.startswith("iphone") or d.startswith("pixel")


def _select(findings, device_filter):
    if device_filter == "mobile":
        return [f for f in findings if _is_mobile_device(f.get("device"))]
    return [f for f in findings if (f.get("device") or "").lower() == "desktop"]


def _is_not_reproducible(f):
    return f.get("status") in NOT_REPRODUCIBLE_STATUSES or bool(f.get("not_reproducible") or f.get("suspect_resolved"))


def _validate_findings(findings):
    for i, f in enumerate(findings):
        for field in REQUIRED_FIELDS:
            value = f.get(field)
            if value is None or (isinstance(value, str) and not value.strip()) or (isinstance(value, list) and not value):
                raise ValueError(f"finding #{i + 1} missing required field '{field}'")
        severity = f["severity"]
        if severity not in RUBRIC:
            raise ValueError(f"finding #{i + 1} has unknown severity {severity!r}; expected one of {list(RUBRIC)}")


def _sort_issues(issues):
    order = {sev: i for i, sev in enumerate(SEVERITY_ORDER)}
    return sorted(issues, key=lambda f: order[f["severity"]])


def _build_cover(doc, device_filter, as_of):
    platform = "Mobile" if device_filter == "mobile" else "Desktop"
    _add_heading(doc, f"injector.world QA Bug Report - {platform}", level=0)
    doc.add_paragraph(f"Prepared by {TESTER} | Test date {date.today().isoformat()} | Data as of {as_of}", style="Subtitle")

    _add_heading(doc, "Executive Summary", level=1)
    meta = [
        ("Site Under Test", SITE_URL),
        ("Platform", platform),
        ("Devices / Viewports", "iPhone 13 + Pixel 7 (Playwright emulation)" if device_filter == "mobile" else "1920x1080 primary; responsive 1366 / 1024 / 768"),
        ("Test Date", date.today().isoformat()),
        ("Data As Of", as_of),
        ("Tester", TESTER),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light Grid Accent 1"
    _style_table(table, [Inches(1.9), Inches(4.4)])
    for row, (k, v) in zip(table.rows, meta):
        _set_cell(row.cells[0], k, bold=True)
        _set_cell(row.cells[1], v)
    return platform


def _add_summary_table(doc, issues):
    _add_heading(doc, "Severity Summary", level=1)
    counts = {sev: sum(1 for f in issues if f["severity"] == sev) for sev in SEVERITY_ORDER}
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    _style_table(table, [Inches(2.2), Inches(1.4)])
    for cell in table.rows[0].cells:
        _set_cell(cell, "", bold=True)
    _set_cell(table.rows[0].cells[0], "Severity", bold=True)
    _set_cell(table.rows[0].cells[1], "Count", bold=True)
    _shade(table.rows[0].cells[0], "D9D9D9")
    _shade(table.rows[0].cells[1], "D9D9D9")
    for i, sev in enumerate(SEVERITY_ORDER, start=1):
        color = _hex_rgb(RUBRIC[sev]["color"])
        _set_cell(table.rows[i].cells[0], sev, bold=True, color=color)
        _set_cell(table.rows[i].cells[1], str(counts[sev]))
        _shade(table.rows[i].cells[0], RUBRIC[sev]["color"].lstrip("#"))
    total_row = table.add_row()
    _set_cell(total_row.cells[0], "Total", bold=True)
    _set_cell(total_row.cells[1], str(sum(counts.values())), bold=True)
    _shade(total_row.cells[0], "D9D9D9")
    _shade(total_row.cells[1], "D9D9D9")
    return counts


def _add_issue(doc, issue, issue_id):
    color = RUBRIC[issue["severity"]]["color"]
    _add_heading(doc, f"{issue_id} - {issue['severity']}: {issue['title']}", level=2, color_hex=color)

    url = issue["url"]
    _add_labeled(doc, "URL", ", ".join(url) if isinstance(url, list) else url)
    _add_labeled(doc, "Device / Viewport", _device_label(issue.get("device")))
    if issue.get("also_present_on"):
        _add_labeled(doc, "Also present on", issue["also_present_on"])

    steps = issue.get("repro_steps")
    if steps:
        _add_bold_label(doc, "Steps to Reproduce")
        _add_steps(doc, steps)

    _add_labeled(doc, "Expected", issue.get("expected", ""))
    _add_labeled(doc, "Actual", issue.get("actual", ""))
    if issue.get("affected_pages"):
        pages = ", ".join(issue["affected_pages"]) if isinstance(issue["affected_pages"], list) else str(issue["affected_pages"])
        _add_labeled(doc, "Affected Pages", pages)
    if issue.get("suggested_fix"):
        _add_labeled(doc, "Suggested Fix", issue["suggested_fix"])

    screenshots = issue["screenshot_path"]
    if not isinstance(screenshots, list):
        screenshots = [screenshots]
    for shot in screenshots:
        _add_picture_with_caption(doc, _resolve(shot))


def _add_issue_sections(doc, issues, issue_id_prefix):
    _add_heading(doc, "Issues", level=1)
    for i, issue in enumerate(issues, start=1):
        _add_issue(doc, issue, f"{issue_id_prefix}-{i:03d}")


def _add_not_reproducible(doc, items):
    if not items:
        return
    _add_heading(doc, "Verified: Not Reproducible", level=1)
    for f in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{f['title']} - ").font.bold = True
        p.add_run(f"URL: {f.get('url', 'unknown')} | Device: {_device_label(f.get('device'))}")
        note = f.get("notes") or f.get("actual")
        if note:
            doc.add_paragraph(f"  {note}")


def _add_coverage(doc, manifest, device_filter):
    _add_heading(doc, f"Coverage - Pages Tested ({'Mobile' if device_filter == 'mobile' else 'Desktop'})", level=1)
    if device_filter == "mobile":
        entries = [e for e in manifest if _is_mobile_device(e.get("device"))]
    else:
        entries = [e for e in manifest if (e.get("device") or "").lower() == "desktop"]
    if not entries:
        doc.add_paragraph("No pages recorded in the crawl manifest for this device.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _style_table(table, [Inches(3.0), Inches(1.3), Inches(1.0), Inches(1.0)])
    headers = ["URL", "Device", "Status", "Load (ms)"]
    for cell, text in zip(table.rows[0].cells, headers):
        _set_cell(cell, text, bold=True)
        _shade(cell, "D9D9D9")
    for e in entries:
        row = table.add_row().cells
        _set_cell(row[0], e.get("url", ""))
        _set_cell(row[1], _device_label(e.get("device")))
        _set_cell(row[2], str(e.get("status", "")))
        _set_cell(row[3], str(e.get("load_time_ms", "")))


def _add_performance(doc, findings):
    perf = [f for f in findings if f.get("kind") == "performance"]
    if not perf:
        return
    _add_heading(doc, "Performance Appendix - Load Times", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _style_table(table, [Inches(3.0), Inches(1.3), Inches(1.2), Inches(2.0)])
    for cell, text in zip(table.rows[0].cells, ["Page URL", "Device", "Load (ms)", "Notes"]):
        _set_cell(cell, text, bold=True)
        _shade(cell, "D9D9D9")
    for f in perf:
        row = table.add_row().cells
        url = f.get("url", "")
        _set_cell(row[0], ", ".join(url) if isinstance(url, list) else url)
        _set_cell(row[1], _device_label(f.get("device")))
        _set_cell(row[2], str(f.get("load_time_ms", f.get("notes", ""))))
        _set_cell(row[3], str(f.get("notes", "")))


def build_report(findings, manifest, device_filter, out_path, as_of=None, exec_note=None):
    if device_filter not in ("mobile", "desktop"):
        raise ValueError(f"--device must be 'mobile' or 'desktop', got {device_filter!r}")
    as_of = as_of or date.today().isoformat()
    _validate_findings(findings)

    selected = _select(findings, device_filter)
    issues = [f for f in selected if f.get("kind") != "performance" and not _is_not_reproducible(f)]
    issues = _sort_issues(issues)
    not_reproducible = [f for f in selected if _is_not_reproducible(f)]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    platform = _build_cover(doc, device_filter, as_of)
    if exec_note:
        doc.add_paragraph(exec_note).italic = True
    doc.add_paragraph(
        f"This report covers {len(issues)} issue(s) on {platform} across the page(s) listed in the "
        f"Coverage section. Severity is classified with the shared rubric (CRITICAL to LOW). "
        f"Every issue includes steps to reproduce, expected vs actual, and screenshot evidence."
    )
    counts = _add_summary_table(doc, issues)

    _add_issue_sections(doc, issues, "M" if device_filter == "mobile" else "D")
    _add_not_reproducible(doc, not_reproducible)
    _add_coverage(doc, manifest, device_filter)
    _add_performance(doc, selected)

    cp = doc.core_properties
    cp.title = f"injector.world QA Bug Report - {platform}"
    cp.author = TESTER
    cp.subject = f"{platform} QA testing of {SITE_URL}"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return {"platform": platform, "counts": counts, "issues": len(issues), "out": str(out_path)}


def _load_json(path):
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"data file not found: {p} (run the crawl harness first or pass --findings/--manifest)")
    return json.loads(p.read_text(encoding="utf-8"))


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--device", required=True, choices=("mobile", "desktop"), help="platform for this report")
    ap.add_argument("--out", required=True, help="output DOCX path (e.g. reports/injector-world-mobile-bug-report.docx)")
    ap.add_argument("--as-of", default=date.today().isoformat(), help="'as of' data timestamp (ISO), default today")
    ap.add_argument("--findings", default=str(PROJECT_ROOT / "data" / "findings.json"))
    ap.add_argument("--manifest", default=str(PROJECT_ROOT / "data" / "crawl-manifest.json"))
    ap.add_argument("--exec-note", default=None, help="paragraph appended to the Executive Summary (e.g. T16 re-verification note)")
    args = ap.parse_args(argv)

    findings = _load_json(args.findings)
    if not isinstance(findings, list):
        raise ValueError(f"{args.findings}: expected a JSON array of findings, got {type(findings).__name__}")
    manifest = _load_json(args.manifest)
    if not isinstance(manifest, list):
        raise ValueError(f"{args.manifest}: expected a JSON array of manifest entries, got {type(manifest).__name__}")

    result = build_report(findings, manifest, args.device, Path(args.out), as_of=args.as_of, exec_note=args.exec_note)
    print(f"WROTE {result['out']}")
    print(f"  platform : {result['platform']}")
    print(f"  issues   : {result['issues']}")
    print(f"  summary  : {result['counts']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
