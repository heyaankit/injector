#!/usr/bin/env python3
"""injector.world DOCX Validation Gate (T15).

Validates the generated QA bug reports (DOCX) end-to-end. Six checks per report:

  1. Reopen        - python-docx parses the file; paragraphs >= 30; tables >= 1
  2. Render        - soffice headless converts the DOCX to a non-empty PDF (exit 0)
  3. Images        - number of word/media/ zip parts >= number of reported issues
  4. Placeholders  - zero occurrences of TBD / [insert / TODO / lorem in the text
                     (legitimate "<input placeholder=...>" a11y text is NOT flagged)
  5. Dedup         - no two issues in the same report share identical title + URL
  6. Manifest      - every issue URL present in data/crawl-manifest.json for the
                     report's device (URLs normalized: strip query + trailing '/').
                     A URL missing from the manifest is allowed ONLY when it is a
                     documented broken-link target (title matches the
                     "Broken navigation: header/footer link does not reach a
                     valid page" finding pattern) - such 404 targets are never
                     crawled and therefore legitimately absent from the manifest.
                     Additionally, every static seed page (first 16 entries of
                     data/seed-urls.json) must be present in the manifest for the
                     report's device.

Strictly read-only against data/ and reports/. Exit code is 0 iff ALL checks pass
for ALL configured reports. The final stdout line is "ALL CHECKS PASSED".

Usage:
    python3 scripts/validate_reports.py                     # both real reports
    python3 scripts/validate_reports.py --desktop PATH      # single report
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOFFICE = "/usr/bin/soffice"
PDF_DIR = Path("/tmp/opencode")

PLACEHOLDER_RE = re.compile(r"(tbd|\[insert|todo|lorem)", re.IGNORECASE)
# Finding titles that document a confirmed-broken (404) link target. Such targets
# are never crawled, so they are legitimately absent from crawl-manifest.json.
BROKEN_TITLE_RE = re.compile(r"broken navigation|does not reach a valid page", re.IGNORECASE)
ISSUE_HEADING_RE = re.compile(r"^(?:[MD])-\d{3} - [A-Z]+: (.+)$")
STATIC_PAGE_COUNT = 16
RENDER_TIMEOUT_S = 120

DEFAULT_REPORTS = {
    "mobile": PROJECT_ROOT / "reports" / "injector-world-mobile-bug-report.docx",
    "desktop": PROJECT_ROOT / "reports" / "injector-world-desktop-bug-report.docx",
}


def load_json(path: Path):
    if not path.exists():
        raise FileNotFoundError(f"data file not found: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def norm_url(u) -> str:
    u = (u or "").strip()
    u = u.split("?")[0].split("#")[0]
    return u.rstrip("/")


def is_mobile_device(device) -> bool:
    d = (device or "").lower()
    return d == "mobile" or "mobile" in d or d.startswith("iphone") or d.startswith("pixel")


def is_not_reproducible(f: dict) -> bool:
    return f.get("status") in ("not_reproducible", "suspect_resolved") or bool(
        f.get("not_reproducible") or f.get("suspect_resolved")
    )


def select_issues(findings, device_filter):
    if device_filter == "mobile":
        sel = [f for f in findings if is_mobile_device(f.get("device"))]
    else:
        sel = [f for f in findings if (f.get("device") or "").lower() == "desktop"]
    return [f for f in sel if f.get("kind") != "performance" and not is_not_reproducible(f)]


def docx_text(doc, docx_path: Path) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    parts.extend(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.S))
    return "\n".join(parts)


def extract_issue_pairs(doc):
    pairs = []
    pending_title = None
    for p in doc.paragraphs:
        m = ISSUE_HEADING_RE.match(p.text.strip())
        if m:
            pending_title = m.group(1).strip()
            continue
        if pending_title is not None and p.text.strip().startswith("URL: "):
            pairs.append((pending_title, p.text.strip()[len("URL: "):].strip()))
            pending_title = None
    return pairs


def check_render(docx_path: Path, device: str):
    if not Path(SOFFICE).exists():
        return False, f"soffice not found at {SOFFICE}"
    profile = f"file:///tmp/opencode/lo-profile-{device}"
    cmd = [
        SOFFICE, "--headless", f"-env:UserInstallation={profile}",
        "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(docx_path),
    ]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=RENDER_TIMEOUT_S)
    except subprocess.TimeoutExpired:
        return False, f"soffice timed out after {RENDER_TIMEOUT_S}s"
    pdf = PDF_DIR / (docx_path.stem + ".pdf")
    size = pdf.stat().st_size if pdf.exists() else 0
    ok = proc.returncode == 0 and size > 0
    detail = f"exit={proc.returncode}, pdf={pdf} ({size} bytes)"
    return ok, detail


def check_manifest(issues, manifest, seed, device_filter):
    desktop_set = {norm_url(e.get("url")) for e in manifest if e.get("device") == "desktop"}
    mobile_set = {norm_url(e.get("url")) for e in manifest
                  if e.get("device") in ("iPhone 13", "Pixel 7")}
    all_manifest = desktop_set | mobile_set
    broken_targets = {norm_url(f.get("url")) for f in issues
                      if isinstance(f.get("url"), str)
                      and BROKEN_TITLE_RE.search(f.get("title", ""))}

    device_set = mobile_set if device_filter == "mobile" else desktop_set

    missing = []
    exempted = []
    for f in issues:
        n = norm_url(f.get("url"))
        if n in device_set:
            continue
        if n in broken_targets and n not in all_manifest:
            exempted.append(n)
            continue
        missing.append((f.get("id"), f.get("url")))

    seed_urls = seed if isinstance(seed, list) else seed.get("urls", [])
    static = seed_urls[:STATIC_PAGE_COUNT]
    missing_static = [u for u in static if norm_url(u) not in device_set]

    ok = not missing and not missing_static
    detail = (
        f"{len(issues)}/{len(issues)} issue URLs verified "
        f"({len(issues) - len(exempted)} in manifest, {len(exempted)} documented broken-link targets); "
        f"{len(static) - len(missing_static)}/{len(static)} static pages present"
    )
    if missing:
        detail += f"; MISSING: {missing}"
    if missing_static:
        detail += f"; STATIC MISSING: {missing_static}"
    return ok, detail


def run_report(device_filter: str, docx_path: Path, findings, manifest, seed) -> tuple[bool, list[str]]:
    issues = select_issues(findings, device_filter)
    expected = len(issues)
    lines = [f"=== Report: {docx_path} ({device_filter}, {expected} expected issue(s)) ==="]
    report_ok = True

    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        return False, lines + [f"[FAIL] 1. Reopen: could not open DOCX with python-docx: {exc}"]
    n_par, n_tab = len(doc.paragraphs), len(doc.tables)
    ok = n_par >= 30 and n_tab >= 1
    report_ok = report_ok and ok
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 1. Reopen: {n_par} paragraphs, {n_tab} tables")

    ok, detail = check_render(docx_path, device_filter)
    report_ok = report_ok and ok
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 2. Render: {detail}")

    with zipfile.ZipFile(docx_path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    ok = len(media) >= expected
    report_ok = report_ok and ok
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 3. Images: {len(media)} media part(s) >= {expected} issue(s)")

    text = docx_text(doc, docx_path)
    hits = PLACEHOLDER_RE.findall(text)
    ok = not hits
    report_ok = report_ok and ok
    detail = f"placeholder token(s) found: {sorted(set(hits))}" if hits else "no TBD / [insert / TODO / lorem tokens"
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 4. Placeholders: {detail}")

    pairs = extract_issue_pairs(doc)
    if len(pairs) == 0 and expected > 0:
        lines.append(f"[FAIL] 5. Dedup: could not parse issue blocks (0 parsed, expected {expected})")
        return False, lines
    by_key = {}
    for title, url in pairs:
        by_key.setdefault((title.lower().strip(), url.lower().strip()), []).append((title, url))
    dups = [k for k, v in by_key.items() if len(v) > 1]
    ok = not dups
    report_ok = report_ok and ok
    detail = f"{len(dups)} duplicate title+URL pair(s): {dups}" if dups else f"{len(pairs)} issue(s), 0 duplicate title+URL"
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 5. Dedup: {detail}")

    ok, detail = check_manifest(issues, manifest, seed, device_filter)
    report_ok = report_ok and ok
    lines.append(f"[{'PASS' if ok else 'FAIL'}] 6. Manifest: {detail}")

    return report_ok, lines


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--mobile", default=None, help="mobile report DOCX path (default: reports/...-mobile...docx)")
    ap.add_argument("--desktop", default=None, help="desktop report DOCX path (default: reports/...-desktop...docx)")
    ap.add_argument("--findings", default=str(PROJECT_ROOT / "data" / "findings.json"))
    ap.add_argument("--manifest", default=str(PROJECT_ROOT / "data" / "crawl-manifest.json"))
    ap.add_argument("--seed", default=str(PROJECT_ROOT / "data" / "seed-urls.json"))
    args = ap.parse_args(argv)

    findings = load_json(Path(args.findings))
    if not isinstance(findings, list):
        raise ValueError(f"{args.findings}: expected a JSON array")
    manifest = load_json(Path(args.manifest))
    seed = load_json(Path(args.seed))

    reports = {}
    if args.mobile is not None:
        reports["mobile"] = Path(args.mobile)
    if args.desktop is not None:
        reports["desktop"] = Path(args.desktop)
    if not reports:
        reports = {k: v for k, v in DEFAULT_REPORTS.items()}

    print("==== injector.world DOCX Validation Gate (T15) ====")
    all_ok = True
    for device_filter in ("mobile", "desktop"):
        if device_filter not in reports:
            continue
        path = reports[device_filter]
        if not path.exists():
            print(f"[FAIL] Report file missing: {path}")
            all_ok = False
            continue
        ok, lines = run_report(device_filter, path, findings, manifest, seed)
        all_ok = all_ok and ok
        for line in lines:
            print(line)
        print()

    if all_ok:
        print("ALL CHECKS PASSED")
        return 0
    print("VALIDATION FAILED - see FAIL lines above")
    return 1


if __name__ == "__main__":
    sys.exit(main())
