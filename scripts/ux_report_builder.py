#!/usr/bin/env python3
"""injector.world UI/UX report + proposal builder (python-docx).

Reusable builder for the UX audit deliverables. It supports two modes:

  python3 scripts/ux_report_builder.py --device desktop --out reports/injector-world-ux-desktop-report.docx
  python3 scripts/ux_report_builder.py --device mobile  --out reports/injector-world-ux-mobile-report.docx
  python3 scripts/ux_report_builder.py --proposal --out reports/injector-world-ux-proposal.docx

Each document is compact and Word-copyable: it uses only standard Word styles
(Title, Heading 1/2/3, Normal, Table Grid / Light Grid Accent 1) and embeds no
images. Findings render strictly as ID + title + one-line description plus a
URL / Related-QA tag line. Evidence lives in the findings stores and is listed
in the Evidence Index appendix.

Device mode reads the matching findings store plus data/ux-coverage.json.
Proposal mode reads both findings stores, data/ux-scores.json and
data/ux-priority-dedup.json.
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ux_rubric import UX_TIERS, SEVERITY_ORDER

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SITE_URL = "https://www.injector.world/"

FINDINGS_FILES = {
    "desktop": PROJECT_ROOT / "data" / "ux-desktop-findings.json",
    "mobile": PROJECT_ROOT / "data" / "ux-mobile-findings.json",
}
COVERAGE_FILE = PROJECT_ROOT / "data" / "ux-coverage.json"
SCORES_FILE = PROJECT_ROOT / "data" / "ux-scores.json"
PRIORITY_FILE = PROJECT_ROOT / "data" / "ux-priority-dedup.json"

DEVICE_LABEL = {
    "desktop": "Desktop",
    "iphone-13": "iPhone 13",
    "pixel-7": "Pixel 7",
    "both": "Both",
}

DEVICE_PROFILES = {
    "desktop": "1920x1080 primary viewport; responsive spot-checks at 1366, 1024 and 768",
    "mobile": "iPhone 13 (390x664) and Pixel 7 (412x839) Playwright emulation",
}

GROUP_ORDER = [
    "static",
    "nav_footer",
    "guides",
    "news",
    "brands",
    "services",
    "state_city",
]

TIER_COLOR = {tier: UX_TIERS[tier]["color"] for tier in UX_TIERS}


def _hex_rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell(cell, text, bold=False, color=None, size=None):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _style_table(table, widths):
    table.autofit = False
    for col, width in zip(table.columns, widths):
        col.width = width


def _add_heading(doc, text, level, color_hex=None):
    heading = doc.add_heading(text, level=level)
    if color_hex is not None:
        for run in heading.runs:
            run.font.color.rgb = _hex_rgb(color_hex)
    return heading


def _add_header_row(table, labels, fill="D9D9D9"):
    for cell, text in zip(table.rows[0].cells, labels):
        _set_cell(cell, text, bold=True)
        _shade(cell, fill)


def _new_doc():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    return doc


def _load_json(path):
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"data file not found: {p}")
    return json.loads(p.read_text(encoding="utf-8"))


def _qa_ref_text(finding):
    qa_ref = finding.get("qa_ref")
    if not qa_ref:
        return None
    if isinstance(qa_ref, list):
        return ", ".join(str(r) for r in qa_ref)
    return str(qa_ref)


def _findings_by_id(findings):
    return {f["id"]: f for f in findings}


def _coverage_groups(coverage):
    return Counter(p.get("group", "other") for p in coverage.get("pages", []))


def _rel_path(path_str):
    try:
        return str(Path(path_str).relative_to(PROJECT_ROOT))
    except ValueError:
        return str(path_str)


def _sort_by_severity(findings):
    order = {t: i for i, t in enumerate(SEVERITY_ORDER)}
    return sorted(findings, key=lambda f: order.get(f.get("tier"), 99))


def _tier_counts(findings):
    counts = {t: 0 for t in SEVERITY_ORDER}
    for f in findings:
        t = f.get("tier")
        if t in counts:
            counts[t] += 1
    return counts


def _add_meta_table(doc, device, as_of, n_pages):
    platform = "Desktop" if device == "desktop" else "Mobile"
    meta = [
        ("Site Under Test", SITE_URL),
        ("Engagement", "Complete UI/UX audit of the public website"),
        ("Date", as_of),
        ("Platform", platform),
        ("Scope", f"{n_pages} pages across 17 UX check areas; {DEVICE_PROFILES[device]}"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light Grid Accent 1"
    _style_table(table, [Inches(1.9), Inches(4.4)])
    for row, (k, v) in zip(table.rows, meta):
        _set_cell(row.cells[0], k, bold=True)
        _set_cell(row.cells[1], v)


def _add_executive_overview(doc, device, findings):
    _add_heading(doc, "Executive Overview", level=1)
    counts = _tier_counts(findings)
    total = len(findings)
    if device == "desktop":
        text = (
            f"This report presents the desktop UI/UX audit of injector.world. "
            f"All 67 accessible pages were reviewed against 17 UX check areas and each "
            f"finding was classified into one of four severity tiers. {total} issues were "
            f"identified, including {counts['Critical']} Critical and "
            f"{counts['Most Important']} Most Important, with the most significant friction "
            f"concentrated in site search, conflicting clinic claims, and navigation."
        )
    else:
        text = (
            f"This report presents the mobile UI/UX audit of injector.world on iPhone 13 and "
            f"Pixel 7 emulation. All 67 accessible pages were reviewed against 17 UX check "
            f"areas on both devices. {total} issues were identified, including "
            f"{counts['Critical']} Critical and {counts['Most Important']} Most Important, "
            f"with the most significant friction in the hamburger menu and touch targets."
        )
    doc.add_paragraph(text)


def _add_coverage(doc, device, coverage):
    _add_heading(doc, "Scope and Coverage", level=1)
    groups = _coverage_groups(coverage)
    total = sum(groups.values())
    if device == "desktop":
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _style_table(table, [Inches(3.2), Inches(1.6)])
        _add_header_row(table, ["Page type group", "Pages tested"])
        for group in GROUP_ORDER:
            row = table.add_row().cells
            _set_cell(row[0], group)
            _set_cell(row[1], str(groups.get(group, 0)))
        total_row = table.add_row().cells
        _set_cell(total_row[0], "Total", bold=True)
        _set_cell(total_row[1], str(total), bold=True)
        _shade(total_row[0], "D9D9D9")
        _shade(total_row[1], "D9D9D9")
        doc.add_paragraph(
            f"Responsive widths checked on desktop: 1920x1080 primary, plus spot-checks at "
            f"1366, 1024 and 768."
        )
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _style_table(table, [Inches(2.6), Inches(1.5), Inches(1.9)])
        _add_header_row(table, ["Page type group", "Pages per device", "Records (2 devices)"])
        for group in GROUP_ORDER:
            row = table.add_row().cells
            _set_cell(row[0], group)
            _set_cell(row[1], str(groups.get(group, 0)))
            _set_cell(row[2], str(groups.get(group, 0) * 2))
        total_row = table.add_row().cells
        _set_cell(total_row[0], "Total", bold=True)
        _set_cell(total_row[1], str(total), bold=True)
        _set_cell(total_row[2], str(total * 2), bold=True)
        _shade(total_row[0], "D9D9D9")
        _shade(total_row[1], "D9D9D9")
        _shade(total_row[2], "D9D9D9")
        doc.add_paragraph(
            f"Each of the {total} pages was tested on both iPhone 13 (390x664) and Pixel 7 "
            f"(412x839) emulation, giving {total * 2} page-device records."
        )
    doc.add_paragraph(coverage.get("methodology", ""))


def _add_findings_by_tier(doc, device, findings):
    _add_heading(doc, "Findings by Tier", level=1)
    by_tier = {t: [] for t in SEVERITY_ORDER}
    for f in findings:
        by_tier.setdefault(f.get("tier"), []).append(f)
    for tier in SEVERITY_ORDER:
        tier_findings = by_tier.get(tier, [])
        if not tier_findings:
            continue
        _add_heading(doc, tier, level=2, color_hex=TIER_COLOR[tier])
        for f in tier_findings:
            p = doc.add_paragraph()
            p.add_run(f"{f['id']} {f['title']}").font.bold = True
            doc.add_paragraph(f["description"])
            tag = doc.add_paragraph()
            tag.add_run(f"URL: {f['url']}")
            if device == "mobile":
                tag.add_run(f" | Device: {DEVICE_LABEL.get(f.get('device'), f.get('device'))}")
            ref = _qa_ref_text(f)
            if ref:
                tag.add_run(f" | Related: {ref}")


def _add_checked_no_issues(doc, store):
    _add_heading(doc, "Appendix: Checked, No Issues", level=1)
    checked = store.get("checked_no_issues", {})
    areas = [k for k, v in checked.items() if v]
    if not areas:
        doc.add_paragraph("No UX check areas were recorded as clear of issues.")
        return
    observations = {o.get("area"): o.get("note", "") for o in store.get("observations", [])}
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _style_table(table, [Inches(2.4), Inches(3.9)])
    _add_header_row(table, ["UX check area", "Note"])
    for area in areas:
        row = table.add_row().cells
        _set_cell(row[0], area)
        _set_cell(row[1], observations.get(area, "No issues found."))
    doc.add_paragraph(
        "The areas above were reviewed during the audit and no issues were found in them."
    )


def _add_evidence_index(doc, findings):
    _add_heading(doc, "Appendix: Evidence Index", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _style_table(table, [Inches(1.4), Inches(4.9)])
    _add_header_row(table, ["Issue ID", "Screenshot"])
    for f in findings:
        row = table.add_row().cells
        _set_cell(row[0], f["id"])
        _set_cell(row[1], _rel_path(f.get("screenshot_path", "")))
    doc.add_paragraph(
        "Screenshots are stored in the evidence library and are not embedded in this report "
        "so the document stays compact and copies cleanly into Word."
    )


def _add_methodology(doc, device):
    _add_heading(doc, "Methodology and Limitations", level=1)
    doc.add_paragraph(
        "The audit was performed against the live production website using Playwright "
        "headless Chromium emulation, not a physical device."
    )
    doc.add_paragraph(
        "Findings were recorded by a single evaluator using automated DOM, console, and "
        "network instrumentation."
    )
    doc.add_paragraph(
        "The 67 tested pages cover every route reachable through the sitemap index and the "
        "primary and secondary navigation. Dynamic listings behind the auto sitemap "
        "(12,939 URLs) were representatively sampled rather than exhaustively crawled."
    )
    if device == "desktop":
        doc.add_paragraph(
            "Desktop testing used a 1920x1080 viewport with responsive spot-checks at 1366, "
            "1024 and 768 pixels."
        )
    else:
        doc.add_paragraph(
            "Mobile testing used iPhone 13 and Pixel 7 emulation profiles, including touch "
            "and tap-target checks against the 44px WCAG 2.5.8 minimum."
        )


def build_device_report(device, out_path, as_of=None):
    if device not in FINDINGS_FILES:
        raise ValueError(f"--device must be 'desktop' or 'mobile', got {device!r}")
    as_of = as_of or date.today().isoformat()
    store = _load_json(FINDINGS_FILES[device])
    findings = store.get("findings", [])
    coverage = _load_json(COVERAGE_FILE)

    doc = _new_doc()
    platform = "Desktop" if device == "desktop" else "Mobile"
    _add_heading(doc, f"injector.world UI/UX Report - {platform}", level=0)
    _add_meta_table(doc, device, as_of, len(findings))
    _add_executive_overview(doc, device, findings)
    _add_coverage(doc, device, coverage)
    _add_findings_by_tier(doc, device, findings)
    _add_checked_no_issues(doc, store)
    _add_evidence_index(doc, findings)
    _add_methodology(doc, device)

    cp = doc.core_properties
    cp.title = f"injector.world UI/UX Report - {platform}"
    cp.author = "injector.world UX Audit"
    cp.subject = f"{platform} UI/UX audit of {SITE_URL}"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return {"platform": platform, "findings": len(findings), "out": str(out_path)}


def _add_proposal_exec_summary(doc, desktop_findings, mobile_findings, scores):
    _add_heading(doc, "Executive Summary", level=1)
    d = len(desktop_findings)
    m = len(mobile_findings)
    overall = scores["overall"]["score"]
    doc.add_paragraph(
        f"A complete UI/UX audit of injector.world was performed on desktop and mobile. "
        f"All 67 accessible pages were tested on each platform across 17 UX check areas, "
        f"producing {d} desktop and {m} mobile findings classified into four severity tiers. "
        f"The most significant issues are a site search that returns no results, conflicting "
        f"clinic claims on the homepage, and navigation links that return HTTP 404. The "
        f"estimated overall UX quality score is {overall:.1f} out of 10."
    )


def _add_proposal_methodology(doc):
    _add_heading(doc, "Methodology and Scope", level=2)
    doc.add_paragraph(
        "All pages reachable via the sitemap index and the primary and secondary navigation "
        "(67 pages) were tested on desktop and mobile. Dynamic listings behind the auto "
        "sitemap (12,939 URLs) were representatively sampled rather than exhaustively crawled."
    )
    doc.add_paragraph(
        "Testing used Playwright headless Chromium emulation, not physical devices, and was "
        "performed by a single evaluator using automated DOM, console, and network "
        "instrumentation."
    )


def _add_platform_findings(doc, title, findings, report_name, top_n=6):
    _add_heading(doc, title, level=1)
    counts = _tier_counts(findings)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _style_table(table, [Inches(2.4), Inches(1.4)])
    _add_header_row(table, ["Tier", "Count"])
    for tier in SEVERITY_ORDER:
        row = table.add_row().cells
        _set_cell(row[0], tier, bold=True, color=_hex_rgb(TIER_COLOR[tier]))
        _set_cell(row[1], str(counts[tier]))
    total_row = table.add_row().cells
    _set_cell(total_row[0], "Total", bold=True)
    _set_cell(total_row[1], str(len(findings)), bold=True)
    _shade(total_row[0], "D9D9D9")
    _shade(total_row[1], "D9D9D9")

    _add_heading(doc, "Top Issues", level=2)
    for f in _sort_by_severity(findings)[:top_n]:
        p = doc.add_paragraph()
        p.add_run(f"{f['id']} {f['title']}").font.bold = True
        doc.add_paragraph(f["description"])
    doc.add_paragraph(
        f"The full list of {len(findings)} findings appears in the {report_name}."
    )


def _add_priority_list(doc, priority, by_id):
    _add_heading(doc, "Priority-wise Issue List", level=1)
    entries = sorted(
        priority.get("entries", []),
        key=lambda e: SEVERITY_ORDER.index(e.get("tier")) if e.get("tier") in SEVERITY_ORDER else 99,
    )
    for entry in entries:
        p = doc.add_paragraph()
        p.add_run(f"[{entry.get('devices')}] ").font.bold = True
        p.add_run(entry.get("title", "")).font.bold = True
        first = entry.get("finding_ids", [None])[0]
        desc = by_id.get(first, {}).get("description", "")
        if desc:
            doc.add_paragraph(desc)


def _add_score_section(doc, scores):
    _add_heading(doc, "Overall UX Quality Score", level=1)
    desktop = scores["desktop"]["score"]
    mobile = scores["mobile"]["score"]
    overall = scores["overall"]["score"]
    doc.add_paragraph(f"Desktop UX Quality Score: {desktop:.1f} / 10")
    doc.add_paragraph(f"Mobile UX Quality Score: {mobile:.1f} / 10")
    doc.add_paragraph(f"Overall UX Quality Score: {overall:.1f} / 10")
    doc.add_paragraph(
        "The score uses a ceiling model: the most severe tier present sets the ceiling "
        "(Critical 5.0, Most Important 7.0, Important 8.5, Normal 9.5) and each additional "
        "issue subtracts a small penalty, clamped to 1.0 to 10.0 and rounded to one decimal."
    )
    doc.add_paragraph(
        f"Desktop inputs: {scores['desktop']['tier_counts']}. "
        f"Mobile inputs: {scores['mobile']['tier_counts']}. "
        f"Overall is the average of the two device scores."
    )
    doc.add_paragraph(
        "This is an estimated score based on expert heuristic review of the tested surface."
    )


def _add_conclusion(doc, overall):
    _add_heading(doc, "Conclusion", level=1)
    doc.add_paragraph(
        f"injector.world is a functional directory with a broadly usable layout on both "
        f"desktop and mobile, but the experience is held back by a small number of "
        f"high-impact issues. The most serious are a search that returns no results, "
        f"conflicting clinic claims on the homepage, and navigation links that return "
        f"HTTP 404, all of which directly affect a visitor's ability to find an injector."
    )
    doc.add_paragraph(
        f"Taken together these issues place the estimated overall UX quality at "
        f"{overall:.1f} out of 10. This is an assessment of the current state of the site "
        f"and is not an action plan."
    )


def build_proposal(out_path, as_of=None):
    as_of = as_of or date.today().isoformat()
    desktop_store = _load_json(FINDINGS_FILES["desktop"])
    mobile_store = _load_json(FINDINGS_FILES["mobile"])
    scores = _load_json(SCORES_FILE)
    priority = _load_json(PRIORITY_FILE)
    desktop_findings = desktop_store.get("findings", [])
    mobile_findings = mobile_store.get("findings", [])
    by_id = {
        **_findings_by_id(desktop_findings),
        **_findings_by_id(mobile_findings),
    }

    doc = _new_doc()
    _add_heading(doc, "injector.world - UI/UX Audit Proposal", level=0)
    doc.add_paragraph(f"Prepared {as_of}", style="Subtitle")
    _add_proposal_exec_summary(doc, desktop_findings, mobile_findings, scores)
    _add_proposal_methodology(doc)
    _add_platform_findings(doc, "Desktop Findings", desktop_findings, "Desktop UI/UX Report")
    _add_platform_findings(doc, "Mobile Findings", mobile_findings, "Mobile UI/UX Report")
    _add_priority_list(doc, priority, by_id)
    _add_score_section(doc, scores)
    _add_conclusion(doc, scores["overall"]["score"])

    cp = doc.core_properties
    cp.title = "injector.world - UI/UX Audit Proposal"
    cp.author = "injector.world UX Audit"
    cp.subject = f"UI/UX audit proposal for {SITE_URL}"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return {"out": str(out_path)}


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--device", choices=("desktop", "mobile"),
                    help="device report to build (desktop or mobile)")
    ap.add_argument("--proposal", action="store_true",
                    help="build the client proposal instead of a device report")
    ap.add_argument("--out", required=True, help="output DOCX path")
    ap.add_argument("--as-of", default=None, help="report date (ISO), default today")
    args = ap.parse_args(argv)

    if args.proposal and args.device:
        ap.error("use either --proposal or --device, not both")
    if not args.proposal and not args.device:
        ap.error("provide --device desktop|mobile or --proposal")

    out = Path(args.out)
    if args.proposal:
        result = build_proposal(out, as_of=args.as_of)
        print(f"WROTE {result['out']}")
    else:
        result = build_device_report(args.device, out, as_of=args.as_of)
        print(f"WROTE {result['out']}")
        print(f"  platform : {result['platform']}")
        print(f"  findings : {result['findings']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())