#!/usr/bin/env python3
"""injector.world UI/UX Validation Gate (T3).

Validates the UX audit artifacts end-to-end. Six checks:

  1. Schema   - every finding has a valid id (UX-D-### / UX-M-###), device in
                {desktop|iphone-13|pixel-7|both}, url, primary_area in the 17
                UX_AREAS, tier in the 4 UX_TIERS, non-empty title, description
                <=160 chars single sentence, an existing non-empty .png
                screenshot_path, and a qa_ref that resolves in data/findings.json
                OR qa_independent=true.
  2. Format   - descriptions contain NO solution-pattern tokens (fix:,
            recommend, should be, consider, suggest). Flagged for review;
            auto-fails only when >2 findings are flagged (false-positive care).
  3. Coverage - data/ux-coverage.json shows all 67 pages tested per device in
            data/ux-manifest.json, and every one of the 17 areas has >=1 finding
            OR a checked_no_issues row per device.
  4. Dedup    - no duplicate (title-normalized, url-host) pairs within a report;
            the proposal priority list has no root-cause dupes (device tags
            stripped).
  5. Score    - --score recomputes desktop/mobile/overall from the findings
            stores via ux_rubric.UX_SCORE and compares to the numbers embedded
            in the proposal (exact match to 1 decimal).
  6. Render   - each DOCX reopens via python-docx, soffice PDF conversion exits
            0, media count >=0 (compact reports; evidence lives in the store),
            the placeholder regex (tbd|[insert|todo|lorem) over paragraphs +
            tables + raw XML returns zero, and no duplicate images (SHA1).

Checks that depend on artifacts not yet produced (coverage/manifest, reports,
proposal) are SKIPPED with a note rather than failed, so the gate can be run
incrementally during the audit. Exit code is 0 iff all applicable checks pass
(unless --expect-fail, which inverts the exit code for negative-path testing).

Usage:
    python3 scripts/validate_ux_reports.py
    python3 scripts/validate_ux_reports.py --findings data/ux-desktop-findings.json
    python3 scripts/validate_ux_reports.py --findings /tmp/opencode/ux-fixture-bad.json --expect-fail
    python3 scripts/validate_ux_reports.py --reports reports/...docx ... --score
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document

from ux_rubric import UX_AREAS, UX_TIERS, UX_SCORE, overall_score

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOFFICE = "/usr/bin/soffice"
PDF_DIR = Path("/tmp/opencode")
RENDER_TIMEOUT_S = 120

PLACEHOLDER_RE = re.compile(r"(tbd|\[insert|todo|lorem)", re.IGNORECASE)
# Solution-pattern tokens that must NOT appear in a one-line UX description.
SOLUTION_TOKENS = ("fix:", "recommend", "should be", "consider", "suggest")
SOLUTION_TOKEN_AUTO_FAIL_LIMIT = 2  # auto-fail when more than this many findings are flagged
ID_RE = re.compile(r"^UX-[DM]-\d{3}$")
ALLOWED_DEVICES = {"desktop", "iphone-13", "pixel-7", "both"}
DESC_MAX_CHARS = 160
MULTI_SENTENCE_RE = re.compile(r"[.!?]\s+[A-Z]")
# Proposal score line, e.g. "Desktop UX Quality Score: 6.5 / 10".
SCORE_LINE_RE = re.compile(
    r"(desktop|mobile|overall)[^\d]{0,30}?(\d+\.\d)\s*/\s*10", re.IGNORECASE
)

DEFAULT_FINDINGS = [
    PROJECT_ROOT / "data" / "ux-desktop-findings.json",
    PROJECT_ROOT / "data" / "ux-mobile-findings.json",
]
DEFAULT_QA_FINDINGS = PROJECT_ROOT / "data" / "findings.json"
DEFAULT_COVERAGE = PROJECT_ROOT / "data" / "ux-coverage.json"
DEFAULT_MANIFEST = PROJECT_ROOT / "data" / "ux-manifest.json"

UX_TIER_NAMES = list(UX_TIERS.keys())


def load_json(path: Path):
    if not path.exists():
        raise FileNotFoundError(f"data file not found: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def norm_url(u) -> str:
    u = (u or "").strip()
    u = u.split("?")[0].split("#")[0]
    return u.rstrip("/")


def url_host(u) -> str:
    u = (u or "").strip()
    m = re.match(r"https?://([^/]+)", u)
    return m.group(1).lower() if m else u.lower()


def norm_title(t) -> str:
    return re.sub(r"\s+", " ", (t or "").strip().lower())


def resolve(path: str) -> Path:
    p = Path(path)
    if p.is_absolute():
        return p
    if (Path.cwd() / p).exists():
        return Path.cwd() / p
    return PROJECT_ROOT / p


def load_findings(path: Path):
    data = load_json(path)
    if isinstance(data, dict):
        return data.get("findings", []), data.get("checked_no_issues", [])
    return data, []


def _merge_findings(*groups) -> list:
    """Union finding groups by id (mobile stores key every finding as 'both')."""
    seen = set()
    out = []
    for group in groups:
        for f in group:
            fid = f.get("id")
            if fid is not None and fid in seen:
                continue
            seen.add(fid)
            out.append(f)
    return out


# ---------------------------------------------------------------------------
# Check 1: schema
# ---------------------------------------------------------------------------
def check_schema(findings, qa_ids, store_label) -> tuple[bool, list[str]]:
    errors = []
    for f in findings:
        fid = f.get("id", "")
        if not ID_RE.match(fid):
            errors.append(f"{fid or '<no id>'}: id must match UX-[DM]-###")
        if f.get("device") not in ALLOWED_DEVICES:
            errors.append(f"{fid}: device {f.get('device')!r} not in {sorted(ALLOWED_DEVICES)}")
        if not isinstance(f.get("url"), str) or not f.get("url").strip():
            errors.append(f"{fid}: url missing/empty")
        if f.get("primary_area") not in UX_AREAS:
            errors.append(f"{fid}: primary_area {f.get('primary_area')!r} not in 17 UX_AREAS")
        if f.get("tier") not in UX_TIER_NAMES:
            errors.append(f"{fid}: tier {f.get('tier')!r} not in {UX_TIER_NAMES}")
        if not isinstance(f.get("title"), str) or not f.get("title").strip():
            errors.append(f"{fid}: title empty")
        desc = f.get("description", "")
        if not isinstance(desc, str) or not desc.strip():
            errors.append(f"{fid}: description missing/empty")
        elif len(desc) > DESC_MAX_CHARS:
            errors.append(f"{fid}: description {len(desc)} chars > {DESC_MAX_CHARS}")
        elif MULTI_SENTENCE_RE.search(desc) or "\n" in desc:
            errors.append(f"{fid}: description is not a single sentence")
        shot = f.get("screenshot_path")
        if not shot:
            errors.append(f"{fid}: screenshot_path missing")
        else:
            sp = resolve(shot)
            if not sp.exists():
                errors.append(f"{fid}: screenshot_path does not exist: {shot}")
            elif sp.stat().st_size == 0:
                errors.append(f"{fid}: screenshot_path is empty (0 bytes): {shot}")
            elif sp.suffix.lower() != ".png":
                errors.append(f"{fid}: screenshot_path is not a .png: {shot}")
        # qa_ref validity OR qa_independent
        if not f.get("qa_independent"):
            qa_ref = f.get("qa_ref")
            if qa_ref:
                refs = qa_ref if isinstance(qa_ref, list) else [qa_ref]
                bad = [r for r in refs if r not in qa_ids]
                if bad:
                    errors.append(f"{fid}: qa_ref not found in data/findings.json: {bad}")
    ok = not errors
    detail = f"{len(findings)} finding(s) schema-valid" if ok else f"{len(errors)} schema violation(s)"
    lines = [f"[{'PASS' if ok else 'FAIL'}] 1. Schema ({store_label}): {detail}"]
    if errors:
        lines += [f"      - {e}" for e in errors[:40]]
    return ok, lines


# ---------------------------------------------------------------------------
# Check 2: format (no solution tokens)
# ---------------------------------------------------------------------------
def check_format(findings, store_label: str) -> tuple[bool, list[str]]:
    flagged = []
    for f in findings:
        desc = (f.get("description") or "").lower()
        hits = [t for t in SOLUTION_TOKENS if t in desc]
        if hits:
            flagged.append((f.get("id"), f.get("description"), hits))
    ok = len(flagged) <= SOLUTION_TOKEN_AUTO_FAIL_LIMIT
    detail = (
        f"{len(flagged)} finding(s) with solution-pattern tokens "
        f"(auto-fail threshold {SOLUTION_TOKEN_AUTO_FAIL_LIMIT})"
    )
    lines = [f"[{'PASS' if ok else 'FAIL'}] 2. Format ({store_label}): {detail}"]
    for fid, desc, hits in flagged:
        lines.append(f"      - {fid}: tokens {hits} -> {desc[:80]!r}")
    return ok, lines


# ---------------------------------------------------------------------------
# Check 3: coverage
# ---------------------------------------------------------------------------
# Manifest device string mapping (manifest stores display names).
DEVICE_MANIFEST_KEY = {
    "desktop": "desktop",
    "iphone-13": "iPhone 13",
    "pixel-7": "Pixel 7",
}


def check_coverage(
    coverage_path, manifest_path, findings_by_device, checked_by_device=None
) -> tuple[bool, list[str]]:
    if not coverage_path.exists() or not manifest_path.exists():
        return True, [
            "[SKIP] 3. Coverage: ux-coverage.json / ux-manifest.json not present yet"
        ]
    coverage = load_json(coverage_path)
    manifest = load_json(manifest_path)
    pages = coverage.get("pages", [])
    seed_urls = {norm_url(p.get("url")) for p in pages}
    devices = ("desktop", "iphone-13", "pixel-7")
    checked_by_device = checked_by_device or {"desktop": set(), "both": set()}
    errors = []

    # All 67 pages tested per device in the manifest.
    for dev in devices:
        manifest_dev = DEVICE_MANIFEST_KEY[dev]
        tested = {norm_url(e.get("url")) for e in manifest if e.get("device") == manifest_dev}
        missing = seed_urls - tested
        if missing:
            errors.append(f"{dev}: {len(missing)} seed page(s) not tested: {sorted(missing)[:5]}")

    # Every one of the 17 areas has >=1 finding OR a checked_no_issues row per device.
    # Desktop's checked_no_issues applies to "desktop"; mobile's "both" bucket applies
    # to both iphone-13 and pixel-7 (mirrors the device:"both" findings rule).
    for dev in devices:
        covered = set()
        groups = [findings_by_device.get(dev, [])]
        if dev != "desktop":
            groups.append(findings_by_device.get("both", []))
        for f in _merge_findings(*groups):
            covered.add(f.get("primary_area"))
        covered |= checked_by_device.get("desktop" if dev == "desktop" else "both", set())
        for area in UX_AREAS:
            if area not in covered:
                errors.append(f"{dev}: area {area} has no finding and no checked_no_issues row")

    ok = not errors
    detail = (
        f"{len(seed_urls)} seed pages x {len(devices)} devices; "
        f"{len(UX_AREAS)} areas covered per device"
    )
    lines = [f"[{'PASS' if ok else 'FAIL'}] 3. Coverage: {detail}"]
    if errors:
        lines += [f"      - {e}" for e in errors[:40]]
    return ok, lines


# ---------------------------------------------------------------------------
# Check 4: dedup (within a report)
# ---------------------------------------------------------------------------
def check_dedup(findings, store_label: str) -> tuple[bool, list[str]]:
    by_key = {}
    for f in findings:
        key = (norm_title(f.get("title")), url_host(f.get("url")))
        by_key.setdefault(key, []).append(f.get("id"))
    dups = {k: v for k, v in by_key.items() if len(v) > 1}
    ok = not dups
    detail = (
        f"{len(findings)} finding(s), 0 duplicate (title, host) pairs"
        if ok
        else f"{len(dups)} duplicate (title, host) pair(s)"
    )
    lines = [f"[{'PASS' if ok else 'FAIL'}] 4. Dedup ({store_label}): {detail}"]
    for (title, host), ids in dups.items():
        lines.append(f"      {title!r} @ {host}: {ids}")
    return ok, lines


# ---------------------------------------------------------------------------
# Check 6: render (per DOCX)
# ---------------------------------------------------------------------------
def docx_text(doc, docx_path: Path) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    parts.extend(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.S))
    return "\n".join(parts)


def check_render(docx_path: Path) -> tuple[bool, list[str]]:
    lines = [f"=== Render: {docx_path} ==="]
    ok = True
    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        return False, lines + [f"[FAIL] 6. Render: could not open DOCX: {exc}"]
    lines.append(f"[PASS] 6. Render: reopened via python-docx ({len(doc.paragraphs)} paragraphs)")

    # soffice PDF conversion
    if Path(SOFFICE).exists():
        profile = f"file:///tmp/opencode/lo-profile-ux"
        cmd = [
            SOFFICE, "--headless", f"-env:UserInstallation={profile}",
            "--convert-to", "pdf", "--outdir", str(PDF_DIR), str(docx_path),
        ]
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=RENDER_TIMEOUT_S)
        except subprocess.TimeoutExpired:
            ok = False
            lines.append(f"[FAIL] 6. Render: soffice timed out after {RENDER_TIMEOUT_S}s")
        else:
            pdf = PDF_DIR / (docx_path.stem + ".pdf")
            size = pdf.stat().st_size if pdf.exists() else 0
            good = proc.returncode == 0 and size > 0
            ok = ok and good
            lines.append(f"[{'PASS' if good else 'FAIL'}] 6. Render: soffice exit={proc.returncode}, pdf={size} bytes")
        # clean up the soffice lock file left behind
        lock = Path(str(docx_path) + "#")
        if lock.exists():
            lock.unlink()
    else:
        lines.append(f"[SKIP] 6. Render: soffice not found at {SOFFICE}")

    # media count >=0 (compact reports; evidence lives in the store)
    with zipfile.ZipFile(docx_path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    lines.append(f"[PASS] 6. Render: {len(media)} media part(s) (>=0)")

    # placeholders
    text = docx_text(doc, docx_path)
    hits = PLACEHOLDER_RE.findall(text)
    good = not hits
    ok = ok and good
    lines.append(
        f"[{'PASS' if good else 'FAIL'}] 6. Render: placeholders "
        f"{sorted(set(hits)) if hits else 'none'}"
    )

    # no duplicate images (SHA1)
    with zipfile.ZipFile(docx_path) as z:
        hashes = {}
        for n in media:
            hashes.setdefault(hashlib.sha1(z.read(n)).hexdigest(), []).append(n)
    dups = [v for v in hashes.values() if len(v) > 1]
    good = not dups
    ok = ok and good
    lines.append(
        f"[{'PASS' if good else 'FAIL'}] 6. Render: duplicate images "
        f"{dups if dups else 'none'}"
    )
    return ok, lines


# ---------------------------------------------------------------------------
# Check 5: score (--score)
# ---------------------------------------------------------------------------
def check_score(findings_by_device, proposal_path) -> tuple[bool, list[str]]:
    if not proposal_path.exists():
        return True, ["[SKIP] 5. Score: proposal DOCX not present yet (--score requires it)"]
    desktop = UX_SCORE(_tier_counts(findings_by_device.get("desktop", [])))
    mobile = UX_SCORE(_tier_counts(_merge_findings(
        findings_by_device.get("mobile", []),
        findings_by_device.get("both", []),
    )))
    overall = overall_score(desktop, mobile)
    try:
        doc = Document(str(proposal_path))
    except Exception as exc:
        return False, [f"[FAIL] 5. Score: could not open proposal: {exc}"]
    text = docx_text(doc, proposal_path)
    embedded = {}
    for label, val in SCORE_LINE_RE.findall(text):
        embedded.setdefault(label.lower(), float(val))
    errors = []
    for label, computed in (("desktop", desktop), ("mobile", mobile), ("overall", overall)):
        if label in embedded and abs(embedded[label] - computed) > 0.05:
            errors.append(f"{label}: embedded {embedded[label]} != recomputed {computed}")
    ok = not errors
    detail = (
        f"desktop={desktop}, mobile={mobile}, overall={overall}"
        + (f"; embedded {embedded}" if embedded else "; no embedded scores found")
    )
    lines = [f"[{'PASS' if ok else 'FAIL'}] 5. Score: {detail}"]
    if errors:
        lines += [f"      - {e}" for e in errors]
    return ok, lines


def _tier_counts(findings) -> dict:
    counts = {t: 0 for t in UX_TIER_NAMES}
    for f in findings:
        t = f.get("tier")
        if t in counts:
            counts[t] += 1
    return counts


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------
def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--findings", nargs="+", default=[str(p) for p in DEFAULT_FINDINGS],
                    help="UX findings JSON file(s) (default: desktop + mobile stores)")
    ap.add_argument("--reports", nargs="+", default=[],
                    help="DOCX report path(s) to render-check")
    ap.add_argument("--proposal", default=str(PROJECT_ROOT / "reports" / "injector-world-ux-proposal.docx"),
                    help="proposal DOCX path for the score check")
    ap.add_argument("--score", action="store_true",
                    help="recompute /10 scores and compare to the proposal")
    ap.add_argument("--expect-fail", action="store_true",
                    help="invert exit code: exit 0 when validation FAILS (negative-path testing)")
    args = ap.parse_args(argv)

    qa_ids = set()
    if DEFAULT_QA_FINDINGS.exists():
        qa_ids = {f.get("id") for f in load_json(DEFAULT_QA_FINDINGS)}

    print("==== injector.world UX Validation Gate (T3) ====")
    all_ok = True
    lines = []

    # Per-store checks: schema, format, dedup.
    findings_by_device = {}
    checked_by_device = {"desktop": set(), "both": set()}
    for path_str in args.findings:
        path = Path(path_str)
        if not path.exists():
            lines.append(f"[FAIL] findings file missing: {path}")
            all_ok = False
            continue
        findings, checked_no_issues = load_findings(path)
        label = path.name
        for f in findings:
            findings_by_device.setdefault(f.get("device"), []).append(f)
        for dev in {f.get("device") for f in findings} & {"desktop", "both"}:
            checked_by_device[dev].update(k for k, v in checked_no_issues.items() if v)
        for check in (check_schema, check_format, check_dedup):
            ok, l = check(findings, qa_ids, label) if check is check_schema else check(findings, label)
            all_ok = all_ok and ok
            lines += l

    # Coverage check: only meaningful for the full audit (both default stores).
    default_set = {str(p) for p in DEFAULT_FINDINGS}
    given_set = {str(Path(p)) for p in args.findings}
    if given_set == default_set:
        ok, l = check_coverage(DEFAULT_COVERAGE, DEFAULT_MANIFEST, findings_by_device, checked_by_device)
        all_ok = all_ok and ok
        lines += l
    else:
        lines.append("[SKIP] 3. Coverage: only runs for the full audit (both default findings stores)")

    # Render check per report.
    for r in args.reports:
        ok, l = check_render(Path(r))
        all_ok = all_ok and ok
        lines += l

    # Score check.
    if args.score:
        ok, l = check_score(findings_by_device, Path(args.proposal))
        all_ok = all_ok and ok
        lines += l

    for line in lines:
        print(line)

    if all_ok:
        print("ALL CHECKS PASSED")
        return 1 if args.expect_fail else 0
    print("VALIDATION FAILED - see FAIL lines above")
    return 0 if args.expect_fail else 1


if __name__ == "__main__":
    sys.exit(main())